import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document

# Matn tahrirlovchi oynasi
def yangi_fayl():
    text_area.delete(1.0, tk.END)

def fayl_och():
    fayl = filedialog.askopenfilename(defaultextension=".txt",
                                      filetypes=[("All Files", "*.*"),
                                                 ("Text Documents", "*.txt"),
                                                 ("Word Documents", "*.docx")])
    if fayl:
        if fayl.endswith(".docx"):
            doc = Document(fayl)
            text_area.delete(1.0, tk.END)
            for paragraph in doc.paragraphs:
                text_area.insert(tk.INSERT, paragraph.text + "\n")
        else:
            with open(fayl, "r") as f:
                content = f.read()
                text_area.delete(1.0, tk.END)
                text_area.insert(tk.INSERT, content)

def fayl_saqlash():
    fayl = filedialog.asksaveasfilename(defaultextension=".txt",
                                        filetypes=[("All Files", "*.*"),
                                                   ("Text Documents", "*.txt"),
                                                   ("Word Documents", "*.docx")])
    if fayl:
        content = text_area.get(1.0, tk.END)
        if fayl.endswith(".docx"):
            doc = Document()
            for line in content.split("\n"):
                doc.add_paragraph(line)
            doc.save(fayl)
            messagebox.showinfo("Saqlash", "DOCX fayl muvaffaqiyatli saqlandi!")
        else:
            with open(fayl, "w") as f:
                f.write(content)
                messagebox.showinfo("Saqlash", "Fayl muvaffaqiyatli saqlandi!")

# Oynani yaratish
root = tk.Tk()
root.title("Matn Muharriri (DOCX qo'llab-quvvatlangan)")

# Menyu barini yaratish
menubar = tk.Menu(root)

# Fayl menyusi
fayl_menu = tk.Menu(menubar, tearoff=0)
fayl_menu.add_command(label="Yangi", command=yangi_fayl)
fayl_menu.add_command(label="Ochish", command=fayl_och)
fayl_menu.add_command(label="Saqlash", command=fayl_saqlash)
fayl_menu.add_separator()
fayl_menu.add_command(label="Chiqish", command=root.quit)
menubar.add_cascade(label="Fayl", menu=fayl_menu)

root.config(menu=menubar)

# Matn kiritish maydoni
text_area = tk.Text(root, wrap="word")
text_area.pack(expand=True, fill="both")

# Dastur oynasini boshlash
root.mainloop()
